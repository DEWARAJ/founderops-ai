from __future__ import annotations

from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile, ZipFile, is_zipfile

from docx import Document
from pypdf import PdfReader
from pypdf.errors import PdfReadError

MAX_RESUME_BYTES = 5 * 1024 * 1024
MAX_EXPANDED_DOCX_BYTES = 25 * 1024 * 1024
MAX_PDF_PAGES = 50
SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}


class DocumentError(ValueError):
    pass


def extract_resume_text(filename: str, content: bytes) -> str:
    """Extract text from a bounded, allowlisted résumé document."""
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise DocumentError(f"Unsupported file type. Use one of: {allowed}.")
    if not content:
        raise DocumentError("The uploaded résumé is empty.")
    if len(content) > MAX_RESUME_BYTES:
        raise DocumentError("The uploaded résumé exceeds the 5 MB limit.")

    try:
        if extension == ".txt":
            if b"\x00" in content:
                raise DocumentError("The TXT résumé appears to contain binary data.")
            text = content.decode("utf-8-sig")
        elif extension == ".pdf":
            if not content.lstrip().startswith(b"%PDF-"):
                raise DocumentError("The file does not have a valid PDF signature.")
            reader = PdfReader(BytesIO(content))
            if len(reader.pages) > MAX_PDF_PAGES:
                raise DocumentError(f"PDF résumés are limited to {MAX_PDF_PAGES} pages.")
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        else:
            _validate_docx_archive(content)
            document = Document(BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_cells = [
                cell.text
                for table in document.tables
                for row in table.rows
                for cell in row.cells
            ]
            text = "\n".join([*paragraphs, *table_cells])
    except DocumentError:
        raise
    except (UnicodeDecodeError, OSError, ValueError, KeyError, BadZipFile, PdfReadError) as error:
        raise DocumentError("The résumé could not be parsed as a valid document.") from error

    normalized = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    if len(normalized) < 40:
        raise DocumentError("The résumé does not contain enough extractable text.")
    return normalized


def _validate_docx_archive(content: bytes) -> None:
    buffer = BytesIO(content)
    if not is_zipfile(buffer):
        raise DocumentError("The file does not have a valid DOCX signature.")
    buffer.seek(0)
    with ZipFile(buffer) as archive:
        members = archive.infolist()
        if "word/document.xml" not in {member.filename for member in members}:
            raise DocumentError("The file is not a valid Word document.")
        if sum(member.file_size for member in members) > MAX_EXPANDED_DOCX_BYTES:
            raise DocumentError("The expanded DOCX content exceeds the safe processing limit.")
