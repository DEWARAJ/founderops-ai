from io import BytesIO

import pytest
from docx import Document

from founderops.documents import MAX_RESUME_BYTES, DocumentError, extract_resume_text


def test_extracts_utf8_text_resume() -> None:
    content = b"Engineer with 3+ years building Python and FastAPI production services."

    assert "Python" in extract_resume_text("resume.txt", content)


def test_extracts_docx_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("Engineer building Python and PostgreSQL services for 2+ years.")
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Reduced processing time by 25%."
    buffer = BytesIO()
    document.save(buffer)

    text = extract_resume_text("resume.docx", buffer.getvalue())

    assert "Python and PostgreSQL" in text
    assert "Reduced processing time" in text


@pytest.mark.parametrize("filename", ["resume.exe", "resume.png", "resume.zip"])
def test_rejects_non_document_extensions(filename: str) -> None:
    with pytest.raises(DocumentError, match="Unsupported file type"):
        extract_resume_text(filename, b"not an allowed resume document" * 3)


def test_rejects_documents_without_extractable_content() -> None:
    with pytest.raises(DocumentError, match="enough extractable text"):
        extract_resume_text("resume.txt", b"short")


def test_rejects_oversized_and_disguised_documents() -> None:
    with pytest.raises(DocumentError, match="5 MB"):
        extract_resume_text("resume.txt", b"a" * (MAX_RESUME_BYTES + 1))
    with pytest.raises(DocumentError, match="PDF signature"):
        extract_resume_text("resume.pdf", b"not really a pdf" * 4)
    with pytest.raises(DocumentError, match="DOCX signature"):
        extract_resume_text("resume.docx", b"not really a docx" * 4)


def test_rejects_binary_content_disguised_as_text() -> None:
    with pytest.raises(DocumentError, match="binary data"):
        extract_resume_text("resume.txt", b"Engineer\x00with Python experience" * 3)
